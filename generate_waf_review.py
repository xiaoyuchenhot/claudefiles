"""Generate a sample WAF Review document for an Azure e-commerce workload."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
def set_heading(para, level=1):
    para.style = f'Heading {level}'

def add_colored_cell(cell, text, bg_hex, bold=False, font_size=10):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

PILLAR_COLORS = {
    'Reliability':            'D6E4FF',
    'Security':               'FFE4CC',
    'Cost Optimization':      'D6F5D6',
    'Operational Excellence': 'FFF3CC',
    'Performance Efficiency': 'F5D6FF',
}
RATING_COLORS = {'High': 'FF4444', 'Medium': 'FF9900', 'Low': 'FFD966', 'Met': '4CAF50'}

def rating_color(r):
    for k, v in RATING_COLORS.items():
        if k.lower() in r.lower():
            return v
    return 'FFFFFF'

def add_header_row(table, headers, bg='1F5C99', fg='FFFFFF'):
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(fg)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(11)
section.page_height = Inches(8.5)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(0.75)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Azure Well-Architected Framework Review')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Workload: Contoso E-Commerce Platform')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Review Date: {datetime.date.today().strftime("%B %d, %Y")}   |   '
             f'Prepared by: Cloud Architecture Team   |   Version: 1.0').font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('1. Executive Summary')
set_heading(h, 1)

doc.add_paragraph(
    'This Well-Architected Framework (WAF) review assesses the Contoso E-Commerce Platform, '
    'a customer-facing retail application serving ~500,000 monthly active users across Australia, '
    'Southeast Asia, and North America. The review follows the Azure Architecture Center methodology '
    'and evaluates the workload across all five WAF pillars: Reliability, Security, Cost Optimization, '
    'Operational Excellence, and Performance Efficiency.'
)

# Score table
doc.add_paragraph('Overall Pillar Scores', style='Heading 3')
t = doc.add_table(rows=6, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['WAF Pillar', 'Score (1–5)', 'Rating', 'Priority'])
data = [
    ('Reliability',            '3 / 5', 'Medium Risk',   'High'),
    ('Security',               '4 / 5', 'Low Risk',      'Medium'),
    ('Cost Optimization',      '2 / 5', 'High Risk',     'High'),
    ('Operational Excellence', '3 / 5', 'Medium Risk',   'Medium'),
    ('Performance Efficiency', '3 / 5', 'Medium Risk',   'High'),
]
for i, (pillar, score, risk, pri) in enumerate(data, 1):
    row = t.rows[i]
    add_colored_cell(row.cells[0], pillar, PILLAR_COLORS[pillar])
    row.cells[1].text = score
    add_colored_cell(row.cells[2], risk, rating_color(risk))
    add_colored_cell(row.cells[3], pri,  rating_color(pri))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. WORKLOAD OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('2. Workload Overview')
set_heading(h, 1)

doc.add_paragraph('2.1 Architecture Style', style='Heading 2')
doc.add_paragraph(
    'The platform follows a Microservices architecture deployed on Azure Kubernetes Service (AKS), '
    'fronted by Azure API Management and Azure Front Door for global load balancing. '
    'Core services include: Product Catalog, Order Management, Payment Processing, '
    'User Identity, Inventory, and Notification services.'
)

doc.add_paragraph('2.2 Key Requirements', style='Heading 2')
reqs = [
    ('Availability target', '99.9% (current), target 99.95%'),
    ('Latency (p95)',        '< 300 ms for product browsing; < 500 ms for checkout'),
    ('Peak throughput',     '5,000 requests/sec during sales events'),
    ('Data residency',      'Australia East (primary), Southeast Asia (DR)'),
    ('RTO / RPO',           '4 hours / 1 hour'),
    ('Compliance',          'PCI-DSS (payment), GDPR, Australian Privacy Act'),
]
t2 = doc.add_table(rows=len(reqs)+1, cols=2)
t2.style = 'Table Grid'
add_header_row(t2, ['Requirement', 'Target / Value'])
for i, (req, val) in enumerate(reqs, 1):
    t2.rows[i].cells[0].text = req
    t2.rows[i].cells[1].text = val

doc.add_paragraph()

doc.add_paragraph('2.3 Azure Services in Scope', style='Heading 2')
services = [
    ('Compute',    'AKS (node pools: system + user), Azure Container Apps (background jobs)'),
    ('Data',       'Azure Cosmos DB (orders, catalog), Azure SQL Database (inventory), Redis Cache'),
    ('Messaging',  'Azure Service Bus (order events), Azure Event Grid (notifications)'),
    ('Networking', 'Azure Front Door, API Management, Application Gateway, VNet + Private Endpoints'),
    ('Storage',    'Azure Blob Storage (product images, exports)'),
    ('Identity',   'Microsoft Entra ID, Managed Identity, Key Vault'),
    ('Monitoring', 'Application Insights, Azure Monitor, Log Analytics Workspace'),
]
t3 = doc.add_table(rows=len(services)+1, cols=2)
t3.style = 'Table Grid'
add_header_row(t3, ['Layer', 'Azure Services'])
for i, (layer, svc) in enumerate(services, 1):
    t3.rows[i].cells[0].text = layer
    t3.rows[i].cells[1].text = svc

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. WAF PILLAR ASSESSMENTS
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('3. WAF Pillar Assessments')
set_heading(h, 1)

pillars = [
    {
        'name': 'Reliability',
        'score': '3/5',
        'summary': (
            'The workload has foundational reliability controls in place but lacks full redundancy '
            'at the data layer and has no tested disaster recovery runbook.'
        ),
        'findings': [
            ('AKS node pools are not zone-redundant',              'High',   'Enable availability zones on all AKS node pools'),
            ('No automated failover to DR region tested',          'High',   'Implement and test Azure Site Recovery + Cosmos DB multi-region writes'),
            ('Circuit breaker not implemented on Payment service', 'High',   'Apply Circuit Breaker pattern using Polly library'),
            ('Health endpoints exist but not consumed by LB',      'Medium', 'Configure AKS readiness/liveness probes and Front Door health checks'),
            ('Redis Cache is single instance (no geo-replication)','Medium', 'Enable Azure Cache for Redis geo-replication (Premium tier)'),
            ('RTO/RPO not formally measured',                      'Medium', 'Conduct quarterly DR drills with measured RTO/RPO'),
        ],
        'patterns': ['Circuit Breaker', 'Health Endpoint Monitoring', 'Retry', 'Bulkhead', 'Saga'],
    },
    {
        'name': 'Security',
        'score': '4/5',
        'summary': (
            'Security posture is strong with managed identity and Key Vault in use. '
            'Minor gaps exist in network segmentation and secret rotation automation.'
        ),
        'findings': [
            ('Some services still use connection strings in env vars', 'High',   'Migrate all secrets to Key Vault; use managed identity references'),
            ('No WAF policy on Application Gateway',                   'Medium', 'Enable OWASP 3.2 ruleset on Azure WAF'),
            ('Container images not scanned in CI/CD pipeline',         'Medium', 'Integrate Microsoft Defender for Containers image scanning'),
            ('Key Vault secret rotation not automated',                 'Low',    'Configure Key Vault rotation policies and event-triggered rotation'),
        ],
        'patterns': ['Gatekeeper', 'Federated Identity', 'Valet Key', 'Rate Limiting'],
    },
    {
        'name': 'Cost Optimization',
        'score': '2/5',
        'summary': (
            'Cost is the weakest pillar. AKS node pools are over-provisioned, no reserved instances '
            'are in use, and blob storage lifecycle policies are absent.'
        ),
        'findings': [
            ('AKS node pools always at peak capacity — no autoscaling', 'High',   'Enable KEDA-based autoscaling; configure cluster autoscaler'),
            ('No Azure Reservations or Savings Plans purchased',        'High',   'Purchase 1-year reservations for AKS nodes and Cosmos DB RUs'),
            ('Blob Storage — no lifecycle tiering policy',              'High',   'Move images older than 90 days to Cool, 365 days to Archive tier'),
            ('Dev/test environments run 24/7',                          'Medium', 'Auto-shutdown dev clusters outside business hours with AKS start/stop'),
            ('Cosmos DB provisioned RUs not right-sized',               'Medium', 'Switch to autoscale RUs; monitor normalized RU consumption'),
            ('No cost alerts or budgets configured',                    'Low',    'Set up Azure Cost Management budgets with 80%/100% alerts'),
        ],
        'patterns': ['Compute Resource Consolidation', 'Static Content Hosting'],
    },
    {
        'name': 'Operational Excellence',
        'score': '3/5',
        'summary': (
            'CI/CD pipelines are in place via GitHub Actions. Observability coverage is partial — '
            'distributed tracing is not end-to-end and runbooks are informal.'
        ),
        'findings': [
            ('No distributed tracing across all microservices',    'High',   'Enable OpenTelemetry SDK; configure Application Insights distributed trace'),
            ('Deployment rollback is manual',                      'High',   'Implement blue-green deployments with automated rollback on health failure'),
            ('No structured runbooks for common incidents',        'Medium', 'Author runbooks in Azure Automation for top-5 incident types'),
            ('Infrastructure not fully IaC (manual Azure Portal)', 'Medium', 'Migrate remaining resources to Bicep/Terraform; enforce via policy'),
            ('Log retention is 30 days (compliance requires 90)', 'Medium', 'Update Log Analytics workspace retention to 90 days'),
        ],
        'patterns': ['Health Endpoint Monitoring', 'External Configuration Store', 'Sidecar', 'Ambassador'],
    },
    {
        'name': 'Performance Efficiency',
        'score': '3/5',
        'summary': (
            'Baseline performance targets are defined. Caching is partially used. '
            'The product catalog service exhibits Chatty I/O and Extraneous Fetching antipatterns.'
        ),
        'findings': [
            ('Product Catalog makes N+1 DB queries per request',        'High',   'Refactor to batch queries; apply Cache-Aside with Redis (TTL: 5 min)'),
            ('No CDN for product images',                               'High',   'Enable Azure Front Door CDN for Blob Storage product images'),
            ('No load testing before peak events (e.g., sales)',        'High',   'Implement k6/Azure Load Testing; run pre-event performance baseline'),
            ('Order service uses synchronous payment call (blocking)',  'Medium', 'Refactor to Asynchronous Request-Reply pattern via Service Bus'),
            ('No autoscaling on Container Apps background processors',  'Medium', 'Configure KEDA Service Bus trigger for order processor scaling'),
            ('API responses not compressed (gzip)',                     'Low',    'Enable response compression in API Management and AKS ingress'),
        ],
        'patterns': ['Cache-Aside', 'CQRS', 'Asynchronous Request-Reply', 'Queue-Based Load Leveling', 'CDN/Static Content Hosting'],
    },
]

for pillar in pillars:
    doc.add_paragraph(f'3.{pillars.index(pillar)+1} {pillar["name"]} — Score: {pillar["score"]}', style='Heading 2')
    p = doc.add_paragraph()
    run = p.add_run('Summary: ')
    run.bold = True
    p.add_run(pillar['summary'])

    doc.add_paragraph('Findings & Recommendations:', style='Heading 3')
    t = doc.add_table(rows=len(pillar['findings'])+1, cols=3)
    t.style = 'Table Grid'
    add_header_row(t, ['Finding', 'Severity', 'Recommendation'], bg=list(PILLAR_COLORS.values())[pillars.index(pillar)].replace('FF','AA') if False else '2E6DA4')
    for i, (finding, sev, rec) in enumerate(pillar['findings'], 1):
        t.rows[i].cells[0].text = finding
        add_colored_cell(t.rows[i].cells[1], sev, rating_color(sev))
        t.rows[i].cells[2].text = rec

    p2 = doc.add_paragraph()
    run2 = p2.add_run('Recommended Patterns: ')
    run2.bold = True
    p2.add_run(', '.join(pillar['patterns']))
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CONSOLIDATED RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('4. Consolidated Recommendations & Roadmap')
set_heading(h, 1)

roadmap = [
    # (Priority, Pillar, Action, Effort, Impact)
    ('P1 — Immediate', 'Reliability',            'Enable AKS availability zones',                         'Low',    'High'),
    ('P1 — Immediate', 'Cost Optimization',      'Enable AKS cluster autoscaler + KEDA',                  'Medium', 'High'),
    ('P1 — Immediate', 'Performance',            'Cache Product Catalog with Redis + CDN for images',      'Medium', 'High'),
    ('P1 — Immediate', 'Reliability',            'Implement Circuit Breaker on Payment service',           'Medium', 'High'),
    ('P1 — Immediate', 'Operational Excellence', 'Enable end-to-end distributed tracing (OpenTelemetry)', 'Medium', 'High'),
    ('P2 — 30 days',   'Cost Optimization',      'Purchase Azure Reservations (AKS + Cosmos DB)',          'Low',    'High'),
    ('P2 — 30 days',   'Security',               'Enable WAF policy on Application Gateway (OWASP 3.2)',   'Low',    'High'),
    ('P2 — 30 days',   'Reliability',            'Enable Cosmos DB multi-region writes + DR drill',        'High',   'High'),
    ('P2 — 30 days',   'Performance',            'Refactor Order service to async pattern (Service Bus)',  'High',   'Medium'),
    ('P3 — 90 days',   'Cost Optimization',      'Blob Storage lifecycle tiering policy',                  'Low',    'Medium'),
    ('P3 — 90 days',   'Operational Excellence', 'Migrate all infra to Bicep IaC; enforce via policy',     'High',   'Medium'),
    ('P3 — 90 days',   'Security',               'Automate Key Vault secret rotation',                     'Medium', 'Medium'),
    ('P3 — 90 days',   'Performance',            'Pre-event load testing with Azure Load Testing',          'Medium', 'Medium'),
]

t = doc.add_table(rows=len(roadmap)+1, cols=5)
t.style = 'Table Grid'
add_header_row(t, ['Priority', 'Pillar', 'Action', 'Effort', 'Impact'])
for i, (pri, pillar, action, effort, impact) in enumerate(roadmap, 1):
    row = t.rows[i]
    add_colored_cell(row.cells[0], pri, 'FFD966' if 'Immediate' in pri else ('FFE4CC' if '30' in pri else 'F0F0F0'))
    add_colored_cell(row.cells[1], pillar, PILLAR_COLORS.get(pillar, 'FFFFFF'))
    row.cells[2].text = action
    row.cells[3].text = effort
    add_colored_cell(row.cells[4], impact, 'D6F5D6' if impact == 'High' else 'FFF3CC')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ARCHITECTURE DECISION RECORDS
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('5. Architecture Decision Records (ADRs)')
set_heading(h, 1)

adrs = [
    {
        'id': 'ADR-001',
        'title': 'Adopt Circuit Breaker Pattern for Payment Service',
        'status': 'Accepted',
        'context': (
            'The Payment service is a synchronous dependency for checkout. A downstream payment provider '
            'outage directly causes checkout failures and thread pool exhaustion across the Order service.'
        ),
        'decision': (
            'Implement the Circuit Breaker pattern using the Polly .NET library with a 5-second timeout, '
            '50% failure threshold over 10 requests to open the circuit, and a 30-second half-open probe interval. '
            'Graceful degradation: display "payment temporarily unavailable" and queue the order for retry.'
        ),
        'consequences': (
            'Positive: Prevents cascading failures; improves resilience score. '
            'Negative: Requires code changes to Order service; adds retry/queue infrastructure complexity.'
        ),
    },
    {
        'id': 'ADR-002',
        'title': 'Switch Product Catalog to Read-Through Cache with Redis',
        'status': 'Accepted',
        'context': (
            'Product Catalog service performs N+1 queries to Cosmos DB per page render, '
            'causing high RU consumption and p95 latency of 620 ms — exceeding the 300 ms target.'
        ),
        'decision': (
            'Apply Cache-Aside pattern using Azure Cache for Redis (Standard C2 tier). '
            'Cache product list pages with TTL of 5 minutes; cache individual product details with TTL of 15 minutes. '
            'Invalidate cache on product update events via Event Grid.'
        ),
        'consequences': (
            'Positive: Reduces Cosmos DB RU consumption by ~70%; improves p95 latency to ~80 ms. '
            'Negative: Introduces eventual consistency window of up to 5 minutes for product updates; '
            'adds Redis operational overhead.'
        ),
    },
    {
        'id': 'ADR-003',
        'title': 'Enable AKS Cluster Autoscaler with KEDA for Cost Optimization',
        'status': 'Proposed',
        'context': (
            'AKS node pools are statically sized at peak capacity (20 nodes), resulting in ~65% idle compute '
            'during off-peak hours (10 PM – 8 AM AEST). Monthly compute waste estimated at AUD $4,200.'
        ),
        'decision': (
            'Enable AKS Cluster Autoscaler (min: 5 nodes, max: 25 nodes) on user node pool. '
            'Deploy KEDA with Azure Service Bus triggers to scale Order Processor pods based on queue depth. '
            'Set scale-down delay to 10 minutes to avoid flapping during flash sales.'
        ),
        'consequences': (
            'Positive: Estimated 40–50% reduction in AKS compute costs; auto-scales for peak events. '
            'Negative: Cold-start latency of ~90 seconds when scaling up from minimum; '
            'requires PodDisruptionBudgets to be defined for all services.'
        ),
    },
]

for adr in adrs:
    doc.add_paragraph(f'{adr["id"]}: {adr["title"]}', style='Heading 2')
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    labels = ['Status', 'Context', 'Decision', 'Consequences']
    values = [adr['status'], adr['context'], adr['decision'], adr['consequences']]
    for i, (label, value) in enumerate(zip(labels, values)):
        add_colored_cell(t.rows[i].cells[0], label, 'D6E4FF', bold=True)
        t.rows[i].cells[1].text = value
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ANTIPATTERNS IDENTIFIED
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('6. Performance Antipatterns Identified')
set_heading(h, 1)

antipatterns = [
    ('Chatty I/O',          'Product Catalog — N+1 Cosmos DB queries per page',       'Batch queries; apply Cache-Aside with Redis'),
    ('Extraneous Fetching', 'Order list returns full order object including line items', 'Project only required fields; use CQRS read model'),
    ('Synchronous I/O',     'Payment call blocks Order service thread for up to 8s',   'Async Request-Reply via Service Bus'),
    ('No Caching',          'Homepage product carousel hits DB on every request',       'Redis cache with 5-min TTL; CDN for images'),
    ('Busy Database',       'Reporting queries run on primary Cosmos DB container',     'Separate read replica / materialized view for reports'),
]

t = doc.add_table(rows=len(antipatterns)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Antipattern', 'Observed In', 'Recommended Fix'])
for i, (ap, obs, fix) in enumerate(antipatterns, 1):
    add_colored_cell(t.rows[i].cells[0], ap, 'FFE4CC', bold=True)
    t.rows[i].cells[1].text = obs
    t.rows[i].cells[2].text = fix

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
h = doc.add_paragraph('7. Next Steps')
set_heading(h, 1)

steps = [
    'Schedule P1 items into current sprint backlog (target: 2-week completion)',
    'Assign pillar owners: Engineering Lead (Reliability + Performance), SecOps (Security), FinOps (Cost), SRE (Operational Excellence)',
    'Book follow-up WAF re-assessment in 90 days to validate remediation progress',
    'Conduct DR drill within 30 days — measure actual RTO/RPO against 4h/1h targets',
    'Set up Azure Cost Management budgets and enable Advisor recommendations review (weekly)',
    'Register all ADRs in the team wiki; track decision implementation in project board',
]
for step in steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

doc.add_paragraph()

# Footer note
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'This document was generated using the Cloud Solution Architect skill following '
    'Azure Architecture Center best practices and the Well-Architected Framework methodology.'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output = '/home/user/claudefiles/Contoso_WAF_Review.docx'
doc.save(output)
print(f'Saved: {output}')
